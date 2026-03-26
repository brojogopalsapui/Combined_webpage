from __future__ import annotations

from pathlib import Path
from typing import Iterable, List, Dict, Tuple
import re
import shutil
from html import escape
from calendar import month_name

from bs4 import BeautifulSoup, NavigableString
from docx import Document

TEMPLATE_DOCX_HINTS = {
    "weekly-research-watch-template.docx",
    "template.docx",
    "research-watch-template.docx",
}


def discover_workspace_root(start: Path) -> Path:
    start = start.resolve()
    for candidate in [start, *start.parents]:
        if (candidate / 'AiSecurityResearch').exists() and (candidate / 'brojogopalsapui.github.io').exists():
            return candidate
    raise FileNotFoundError('Could not find workspace root containing both codebases.')


def repo_root(workspace_root: Path, repo_name: str) -> Path:
    path = workspace_root / repo_name
    if not path.exists():
        raise FileNotFoundError(f'Repo not found: {path}')
    return path


def list_site_roots(workspace_root: Path, repo_names: Iterable[str], include_papers_articles: bool = True) -> List[Path]:
    roots: List[Path] = []
    for repo_name in repo_names:
        root = repo_root(workspace_root, repo_name)
        if (root / 'index.html').exists():
            roots.append(root)
        nested = root / 'papers_articles'
        if include_papers_articles and (nested / 'index.html').exists():
            roots.append(nested)
    return roots


def clean_text(value) -> str:
    if value is None:
        return ''
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def stem_to_month_year(stem: str) -> str:
    m = re.match(r'^(\d{4})-(\d{2})-', stem)
    if not m:
        return stem
    year = int(m.group(1))
    month = int(m.group(2))
    if 1 <= month <= 12:
        return f"{month_name[month]} {year}"
    return stem


def pick_latest_docx(folder: Path) -> Path:
    pattern = re.compile(r'^(\d{4})-(\d{2})-(.+)\.docx$', re.IGNORECASE)
    candidates = []
    for f in folder.glob('*.docx'):
        if pattern.match(f.name) and f.name not in TEMPLATE_DOCX_HINTS:
            candidates.append(f)
    if not candidates:
        raise FileNotFoundError(f'No dated DOCX files found in {folder}. Use names like YYYY-MM-topic-name.docx')
    return max(candidates, key=lambda f: f.stat().st_mtime)


def parse_weekly_docx(path: Path) -> Dict:
    doc = Document(str(path))
    data = {'metadata': {}}

    if doc.tables:
        table = doc.tables[0]
        for row in table.rows[1:]:
            key = clean_text(row.cells[0].text)
            val = clean_text(row.cells[1].text)
            if key:
                data['metadata'][key] = val

    wanted_headings = [
        'Preview',
        'Full Note Paragraph 1',
        'Full Note Paragraph 2',
        'What Is Changing Technically',
        'What Reviewers Should Notice',
        'Current Research Tension',
    ]

    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        heading = clean_text(p.text)
        if heading in wanted_headings:
            j = i + 1
            content = []
            while j < len(paragraphs):
                t = clean_text(paragraphs[j].text)
                if t in wanted_headings or t in {'Metadata', 'How the notebook uses this DOCX'}:
                    break
                if t:
                    content.append(t)
                j += 1
            if heading in {'What Is Changing Technically', 'What Reviewers Should Notice'}:
                data[heading] = content
            else:
                data[heading] = '\n'.join(content)
    return data


def normalize_docx_data(data: Dict, docx_stem: str, full_post_link: str | None = None, force_post_id_from_filename: bool = True) -> List[str]:
    md = data['metadata']
    warnings: List[str] = []

    if force_post_id_from_filename:
        existing_post_id = clean_text(md.get('Post ID'))
        if existing_post_id and existing_post_id != docx_stem:
            warnings.append(f"Overriding DOCX Post ID '{existing_post_id}' with filename-based Post ID '{docx_stem}'.")
        md['Post ID'] = docx_stem
    elif not clean_text(md.get('Post ID')):
        md['Post ID'] = docx_stem

    if full_post_link:
        existing_link = clean_text(md.get('Full Post Link (optional)'))
        if existing_link and existing_link != full_post_link:
            warnings.append(f"Overriding DOCX Full Post Link '{existing_link}' with '{full_post_link}'.")
        md['Full Post Link (optional)'] = full_post_link

    if not clean_text(md.get('Title')):
        md['Title'] = docx_stem.replace('-', ' ').title()
        warnings.append('Title was blank in the DOCX, so a title was inferred from the filename.')

    if not clean_text(md.get('Meta Line')):
        md['Meta Line'] = f'Research Watch • {stem_to_month_year(docx_stem)}'
        warnings.append('Meta Line was blank in the DOCX, so one was inferred from the filename.')

    if not clean_text(data.get('Preview')) and clean_text(md.get('Preview')):
        data['Preview'] = clean_text(md.get('Preview'))

    return warnings


def add_link_html(links: List[str], href: str, label: str):
    href = clean_text(href)
    label = clean_text(label)
    if not href or not label:
        return
    if href.startswith('http://') or href.startswith('https://'):
        links.append(f'<a href="{escape(href, quote=True)}" rel="noopener noreferrer" target="_blank">{escape(label)}</a>')
    else:
        links.append(f'<a href="{escape(href, quote=True)}">{escape(label)}</a>')


def build_watch_article(data: Dict, docx_stem: str) -> Tuple[str, str, str, str]:
    md = data['metadata']
    post_id = clean_text(md.get('Post ID')) or docx_stem
    meta_line = clean_text(md.get('Meta Line')) or f'Research Watch • {stem_to_month_year(docx_stem)}'
    title = clean_text(md.get('Title')) or 'Untitled research note'
    preview = clean_text(data.get('Preview')) or clean_text(md.get('Preview'))
    p1 = clean_text(data.get('Full Note Paragraph 1'))
    p2 = clean_text(data.get('Full Note Paragraph 2'))
    tech_list = data.get('What Is Changing Technically', [])
    reviewer_list = data.get('What Reviewers Should Notice', [])
    tension = clean_text(data.get('Current Research Tension'))

    links: List[str] = []
    add_link_html(links, md.get('Full Post Link (optional)'), 'Read full post')
    add_link_html(links, md.get('Related Static Page (optional)'), md.get('Related Static Page Label (optional)'))
    add_link_html(links, md.get('External Link 1 URL (optional)'), md.get('External Link 1 Label (optional)'))
    add_link_html(links, md.get('External Link 2 URL (optional)'), md.get('External Link 2 Label (optional)'))

    tech_items = '\n'.join([f'                      <li>{escape(clean_text(item))}</li>' for item in tech_list if clean_text(item)])
    reviewer_items = '\n'.join([f'                      <li>{escape(clean_text(item))}</li>' for item in reviewer_list if clean_text(item)])
    links_html = ' '.join(links)

    article_html = f'''
<article class="watch-note accordion" id="{escape(post_id, quote=True)}">
  <button aria-expanded="false" class="accordion-trigger" type="button">
    <span class="accordion-meta">{escape(meta_line)}</span>
    <span class="accordion-title">{escape(title)}</span>
    <span class="accordion-preview">
      {escape(preview)}
    </span>
    <span class="accordion-cta">Read full note</span>
    <span aria-hidden="true" class="accordion-icon"></span>
  </button>

  <div aria-hidden="true" class="accordion-panel">
    <div class="accordion-panel-inner">
      <p>
        {escape(p1)}
      </p>

      <p>
        {escape(p2)}
      </p>

      <div class="watch-columns">
        <div class="watch-block">
          <h4>What is changing technically</h4>
          <ul>
{tech_items}
          </ul>
        </div>

        <div class="watch-block">
          <h4>What reviewers should notice</h4>
          <ul>
{reviewer_items}
          </ul>
        </div>
      </div>

      <div class="watch-bottom-note">
        <strong>Current research tension:</strong> {escape(tension)}
      </div>

      <div class="watch-inline-links">
        {links_html}
      </div>
    </div>
  </div>
</article>
'''.strip()

    return post_id, title, preview, article_html


def build_home_watch_card(post_id: str, title: str, preview: str) -> str:
    return f'''
<a class="post-card watch-card" href="ongoing-work.html#{escape(post_id, quote=True)}">
  <span class="tag">Research Watch</span>
  <h3>{escape(title)}</h3>
  <p>
    {escape(preview)}
  </p>
</a>
'''.strip()


def detect_doctype(html_text: str) -> str:
    m = re.match(r'\s*(<!DOCTYPE[^>]+>)', html_text, flags=re.IGNORECASE)
    return m.group(1) if m else '<!DOCTYPE html>'


def update_ongoing_work_html(input_html_path: Path, article_id: str, article_html: str, replace_duplicate: bool = True) -> Path:
    html_text = input_html_path.read_text(encoding='utf-8')
    doctype = detect_doctype(html_text)
    soup = BeautifulSoup(html_text, 'html.parser')

    watch_stack = soup.select_one('div.watch-stack')
    if watch_stack is None:
        raise ValueError(f'Could not find div.watch-stack in {input_html_path}')

    if replace_duplicate:
        existing = watch_stack.find('article', {'id': article_id})
        if existing:
            existing.decompose()

    fragment = BeautifulSoup(article_html, 'html.parser')
    new_article = fragment.find('article')
    first_existing = watch_stack.find('article')
    if first_existing:
        first_existing.insert_before('\n')
        first_existing.insert_before(new_article)
        first_existing.insert_before('\n\n          ')
    else:
        watch_stack.append(new_article)

    final_html = doctype + '\n' + str(soup)
    input_html_path.write_text(final_html, encoding='utf-8')
    return input_html_path


def update_index_html(input_html_path: Path, post_id: str, title: str, preview: str, update_floating: bool = True, update_slider: bool = True, replace_duplicate: bool = True, max_cards: int = 6) -> Path:
    html_text = input_html_path.read_text(encoding='utf-8')
    doctype = detect_doctype(html_text)
    soup = BeautifulSoup(html_text, 'html.parser')

    if update_floating:
        notif_text = soup.select_one('div.floating-notif div.notif-text')
        if notif_text is not None:
            notif_text.clear()
            strong = soup.new_tag('strong')
            strong.string = 'Research Watch:'
            notif_text.append(strong)
            notif_text.append(NavigableString(f' {title}. '))
            link = soup.new_tag('a', href=f'ongoing-work.html#{post_id}')
            link.string = 'Read note →'
            notif_text.append(link)

    if update_slider:
        watch_track = soup.select_one('div#watchTrack')
        if watch_track is not None:
            if replace_duplicate:
                for a in watch_track.select('a.watch-card'):
                    href = a.get('href', '')
                    if href.endswith(f'#{post_id}'):
                        a.decompose()

            card_fragment = BeautifulSoup(build_home_watch_card(post_id, title, preview), 'html.parser')
            new_card = card_fragment.find('a')
            first_card = watch_track.find('a', class_='watch-card')
            if first_card:
                first_card.insert_before('\n')
                first_card.insert_before(new_card)
                first_card.insert_before('\n')
            else:
                watch_track.append(new_card)

            cards = watch_track.select('a.watch-card')
            if max_cards and len(cards) > max_cards:
                for card in cards[max_cards:]:
                    card.decompose()
        else:
            pulse_grid = soup.select_one('div.pulse-grid')
            if pulse_grid is not None:
                if replace_duplicate:
                    for a in pulse_grid.select('a.auto-research-watch'):
                        a.decompose()
                pulse_html = f'''
<a class="pulse-card pulse-card--academic auto-research-watch" href="ongoing-work.html#{escape(post_id, quote=True)}">
  <span class="pulse-label">Research Watch</span>
  <h3>{escape(title)}</h3>
  <p>{escape(preview)}</p>
  <span class="pulse-link">Open note</span>
</a>
'''.strip()
                fragment = BeautifulSoup(pulse_html, 'html.parser')
                new_card = fragment.find('a')
                first_card = pulse_grid.find('a', class_='pulse-card')
                if first_card:
                    first_card.insert_before('\n')
                    first_card.insert_before(new_card)
                    first_card.insert_before('\n')
                else:
                    pulse_grid.append(new_card)

    final_html = doctype + '\n' + str(soup)
    input_html_path.write_text(final_html, encoding='utf-8')
    return input_html_path


def build_full_post_html(data: Dict, docx_stem: str, use_portal_shell: bool = False) -> str:
    md = data['metadata']
    post_id = clean_text(md.get('Post ID')) or docx_stem
    title = clean_text(md.get('Title')) or 'Untitled research note'
    preview = clean_text(data.get('Preview')) or clean_text(md.get('Preview'))
    p1 = clean_text(data.get('Full Note Paragraph 1'))
    p2 = clean_text(data.get('Full Note Paragraph 2'))
    tech_list = data.get('What Is Changing Technically', [])
    reviewer_list = data.get('What Reviewers Should Notice', [])
    tension = clean_text(data.get('Current Research Tension'))
    month_year = stem_to_month_year(docx_stem)

    related_links = []
    add_link_html(related_links, md.get('Related Static Page (optional)'), md.get('Related Static Page Label (optional)'))
    add_link_html(related_links, md.get('External Link 1 URL (optional)'), md.get('External Link 1 Label (optional)'))
    add_link_html(related_links, md.get('External Link 2 URL (optional)'), md.get('External Link 2 Label (optional)'))
    related_links_block = ''
    if related_links:
        related_links_block = '<div class="btn-row" style="margin-top:1rem;">' + ' '.join(related_links) + '</div>'

    tech_items = '\n'.join([f'            <li>{escape(clean_text(item))}</li>' for item in tech_list if clean_text(item)])
    reviewer_items = '\n'.join([f'            <li>{escape(clean_text(item))}</li>' for item in reviewer_list if clean_text(item)])

    if use_portal_shell:
        head_scripts = '<script defer="True" src="../assets/js/site-shell.js"></script><script defer="True" src="../assets/js/portal.js"></script>'
        body_open = f'<body class="subpage portal-page" data-depth="1" data-nav="trending" data-page="{escape(post_id, quote=True)}">\n<div id="site-shell-top"></div><div id="site-shell-header"></div><main>'
        body_close = '</main><div id="site-shell-footer"></div>\n</body>'
    else:
        head_scripts = ''
        body_open = '''<body class="subpage">
<header class="site-header">
<div class="container nav-wrap">
<div class="brand-group"><a aria-label="Brojogopal Sapui Home" class="brand" href="../index.html">B<span>S</span></a><a class="brand-hint" href="https://brojogopalsapui.github.io/AISecurityResearch_Mobileview/" target="_blank" rel="noopener noreferrer">Learning Portal</a></div>
<nav class="nav">
<a href="../index.html">Home</a>
<a href="../about.html">About</a>
<a href="../research.html">Research</a>
<a class="trending-link" href="../ongoing-work.html">Trending Topics</a>
<a href="../publications.html">Resources</a>
<a href="../contact.html">Contact</a>
</nav>
<button aria-expanded="false" aria-label="Toggle menu" class="menu-btn">
<span></span>
<span></span>
<span></span>
</button>
</div>
</header>
<main>'''
        body_close = '''</main>
<footer class="site-footer">
<div class="container footer-grid">
<div>
<h3>Brojogopal Sapui</h3>
<p>AI Security• Hardware Trust • Edge/Physical AI</p>
</div>
<div>
<h4>Main Pages</h4>
<a href="../research.html">Research</a>
<a class="trending-link" href="../ongoing-work.html">Trending Topics</a>
<a href="../publications.html">Resources</a>
</div>
<div>
<h4>Focus</h4>
<p>Cross-layer AI security, trustworthy deployment, hardware-aware defense, and physical intelligence.</p>
</div>
</div>
</footer>
<script src="../assets/js/main.js"></script>
</body>'''

    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<meta content="width=device-width, initial-scale=1.0" name="viewport"/>
<title>{escape(title)} | Brojogopal Sapui</title>
<meta content="{escape(preview)}" name="description"/>
<link href="https://fonts.googleapis.com" rel="preconnect"/>
<link crossorigin="" href="https://fonts.gstatic.com" rel="preconnect"/>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&amp;display=swap" rel="stylesheet"/>
<link href="../assets/css/style.css" rel="stylesheet"/>{head_scripts}</head>
{body_open}
<section class="page-hero">
<div class="container">
<span class="eyebrow">Research Watch • {escape(month_year)}</span>
<h1>{escape(title)}</h1>
<p class="lead">{escape(preview)}</p>
</div>
</section>
<section class="section">
<div class="container split">
<div class="content-card">
<span class="kicker">Overview</span>
<h2>What is changing</h2>
<p>{escape(p1)}</p>
<p>{escape(p2)}</p>
</div>
<div class="content-card">
<span class="kicker">Why it matters</span>
<h2>Research significance</h2>
<ul class="check-list">
{tech_items}
</ul>
</div>
</div>
</section>
<section class="section alt">
<div class="container">
<div class="section-head">
<span class="eyebrow">Discussion</span>
<h2>What reviewers should notice</h2>
<p>
            These review points help separate benchmark-level claims from stronger system-level conclusions.
          </p>
</div>
<div class="content-card">
<ul class="check-list">
{reviewer_items}
</ul>
</div>
</div>
</section>
<section class="section">
<div class="container">
<div class="content-card">
<span class="kicker">Current research tension</span>
<h2>Why this topic matters now</h2>
<p>{escape(tension)}</p>
{related_links_block}
</div>
</div>
</section>
<section class="section alt">
<div class="container">
<div class="cta">
<div>
<span class="eyebrow">Next Step</span>
<h2>Back to ongoing research updates</h2>
<p>
              Return to the running list of research-watch topics and evolving system-level notes.
            </p>
</div>
<div class="cta-actions">
<a class="btn btn-primary" href="../ongoing-work.html#{escape(post_id, quote=True)}">Back to this note</a>
<a class="btn btn-secondary" href="../research.html">Research</a>
</div>
</div>
</div>
</section>
{body_close}
</html>'''.strip()
    return html


def write_full_post_html(output_path: Path, full_post_html: str) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(full_post_html, encoding='utf-8')
    return output_path


def site_uses_portal_shell(site_root: Path) -> bool:
    return (site_root / 'assets' / 'js' / 'site-shell.js').exists() and (site_root / 'assets' / 'js' / 'portal.js').exists()


def sync_weekly_post_to_site(site_root: Path, docx_path: Path, home_slider_max_cards: int = 6, copy_docx_into_site: bool = True, replace_duplicate_post_id: bool = True, update_home_floating: bool = True, update_home_slider: bool = True) -> Dict[str, str]:
    site_root = site_root.resolve()
    docx_stem = docx_path.stem
    data = parse_weekly_docx(docx_path)
    normalize_notes = normalize_docx_data(data, docx_stem, full_post_link=f'posts/{docx_stem}.html')
    post_id, title, preview, article_html = build_watch_article(data, docx_stem)

    if copy_docx_into_site:
        weekly_dir = site_root / 'weekly-inputs'
        weekly_dir.mkdir(parents=True, exist_ok=True)
        shutil.copy2(docx_path, weekly_dir / docx_path.name)

    update_ongoing_work_html(site_root / 'ongoing-work.html', post_id, article_html, replace_duplicate=replace_duplicate_post_id)
    update_index_html(site_root / 'index.html', post_id, title, preview, update_floating=update_home_floating, update_slider=update_home_slider, replace_duplicate=replace_duplicate_post_id, max_cards=home_slider_max_cards)

    post_html = build_full_post_html(data, docx_stem, use_portal_shell=site_uses_portal_shell(site_root))
    post_path = write_full_post_html(site_root / 'posts' / f'{docx_stem}.html', post_html)

    return {
        'site_root': str(site_root),
        'docx_copied_to': str((site_root / 'weekly-inputs' / docx_path.name).resolve()) if copy_docx_into_site else '',
        'ongoing_work': str((site_root / 'ongoing-work.html').resolve()),
        'index_html': str((site_root / 'index.html').resolve()),
        'post_html': str(post_path.resolve()),
        'post_id': post_id,
        'title': title,
        'notes': '; '.join(normalize_notes),
    }


def sync_weekly_post_to_many_sites(site_roots: Iterable[Path], docx_path: Path, **kwargs) -> List[Dict[str, str]]:
    return [sync_weekly_post_to_site(site_root, docx_path, **kwargs) for site_root in site_roots]


def read_manifest(manifest_path: Path) -> List[str]:
    items = []
    for line in manifest_path.read_text(encoding='utf-8').splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith('#'):
            continue
        items.append(stripped)
    return items


def copy_relative_path(source_site_root: Path, target_site_root: Path, relative_path: str) -> Tuple[str, str]:
    src = source_site_root / relative_path
    dst = target_site_root / relative_path
    if not src.exists():
        raise FileNotFoundError(f'Source path not found: {src}')
    dst.parent.mkdir(parents=True, exist_ok=True)
    if src.is_dir():
        if dst.exists():
            shutil.rmtree(dst)
        shutil.copytree(src, dst)
    else:
        shutil.copy2(src, dst)
    return str(src.resolve()), str(dst.resolve())


def sync_article_manifest_to_sites(source_site_root: Path, target_site_roots: Iterable[Path], manifest_entries: Iterable[str]) -> List[Dict[str, str]]:
    results = []
    source_site_root = source_site_root.resolve()
    target_site_roots = [p.resolve() for p in target_site_roots if p.resolve() != source_site_root]
    for rel in manifest_entries:
        for target_site_root in target_site_roots:
            src, dst = copy_relative_path(source_site_root, target_site_root, rel)
            results.append({'relative_path': rel, 'source': src, 'target': dst})
    return results
